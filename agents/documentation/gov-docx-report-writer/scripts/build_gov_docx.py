#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_gov_docx.py — 정부 공문서(제안서) 스타일 .docx 렌더러

설계 의도:
  스타일(글꼴/색상/표제부/절제목 띠/글머리 위계/표/쪽번호/끝.)은 결정적이므로
  이 스크립트가 고정 책임진다. 내용은 JSON 스펙으로 주입한다.
  => 매 문서가 동일한 정부 양식으로 출력된다.

사용법:
  python build_gov_docx.py <spec.json> [output.docx]
  python build_gov_docx.py -            # stdin 으로 JSON 입력
  output 미지정 시 spec 의 meta.output 또는 "<title>.docx" 로 저장

의존성:
  python-docx  (pip install python-docx)

스펙 구조는 sample_spec.json 및 README 참조.
"""

import json
import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "python-docx 가 필요합니다. 다음을 실행하세요:\n"
        "  python -m pip install python-docx\n"
    )
    sys.exit(2)

# ── 스타일 상수 (정부 공문서 스타일 가이드) ──────────────────────────────
FONT_TITLE = "맑은 고딕"   # 제목 · 글머리기호 · 표 머리글
FONT_BODY = "바탕"          # 본문(명조체)

CLR_BLACK = RGBColor(0x00, 0x00, 0x00)
SH_TABLE_HEADER = "E7E6E6"  # 표 머리글 회색
SH_LABEL_COL = "F2F2F2"     # 라벨 열 연회색
SH_SECTION_BAND = "D9D9D9"  # 절 제목 띠 회색
CLR_BORDER = "808080"       # 테두리 회색

# 글머리 위계 마커 (□ 대항목 → ○ 중항목 → - 소항목 → · 세부)
BULLET_MARKERS = {1: "□", 2: "○", 3: "-", 4: "·"}
BULLET_INDENT_CM = {1: 0.0, 2: 0.5, 3: 1.0, 4: 1.5}


# ── 저수준 OOXML 헬퍼 ────────────────────────────────────────────────────
def set_run_font(run, name, size=None, bold=None, color=CLR_BLACK):
    """런에 한글/영문 글꼴을 동시에 지정한다(eastAsia 포함)."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para_shading(paragraph, fill_hex):
    """문단 배경 음영."""
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    ppr.append(shd)


def set_para_border(paragraph, edges):
    """문단 테두리. edges = {'bottom': {'sz':24,'val':'single','color':'000000'}, ...}"""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    for edge, attrs in edges.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), attrs.get("val", "single"))
        el.set(qn("w:sz"), str(attrs.get("sz", 8)))
        el.set(qn("w:space"), str(attrs.get("space", 1)))
        el.set(qn("w:color"), attrs.get("color", "000000"))
        pbdr.append(el)


def set_cell_shading(cell, fill_hex):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcpr.append(shd)


def set_table_borders(table, color=CLR_BORDER, sz=4):
    """표 전체에 얇은 회색 테두리."""
    tbl = table._tbl
    tblpr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblpr.append(borders)


def add_field(run, field_code):
    """런에 필드 코드(예: PAGE)를 삽입한다."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# ── 표제부 ───────────────────────────────────────────────────────────────
def build_header(doc, header):
    """①기관·부서명(굵게)+굵은 가로줄 → ②가운데 큰 제목 → ③'- 부제 -'
       → ④이중 가로줄 → ⑤우측 작성일·작성자"""
    org = header.get("organization", "")
    title = header.get("title", "")
    subtitle = header.get("subtitle", "")
    date = header.get("date", "")
    author = header.get("author", "")

    # ① 기관·부서명 (굵게) + 아래 굵은 가로줄
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(org)
    set_run_font(r, FONT_TITLE, size=12, bold=True)
    set_para_border(p, {"bottom": {"val": "single", "sz": 24, "color": "000000"}})
    p.paragraph_format.space_after = Pt(18)

    # ② 가운데 큰 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, FONT_TITLE, size=22, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

    # ③ '- 부제 -'
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = subtitle if subtitle.strip().startswith("-") else f"- {subtitle} -"
        r = p.add_run(text)
        set_run_font(r, FONT_TITLE, size=13, bold=False)
        p.paragraph_format.space_after = Pt(6)

    # ④ 이중 가로줄
    p = doc.add_paragraph()
    set_para_border(p, {"bottom": {"val": "double", "sz": 6, "color": "000000"}})
    p.paragraph_format.space_after = Pt(4)

    # ⑤ 우측 작성일·작성자
    meta_bits = []
    if date:
        meta_bits.append(date)
    if author:
        meta_bits.append(author)
    if meta_bits:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("    ".join(meta_bits))
        set_run_font(r, FONT_BODY, size=11)
        p.paragraph_format.space_after = Pt(12)


# ── 절 제목 ──────────────────────────────────────────────────────────────
def build_section_heading(doc, number, title):
    """'N. 제목' — 회색 음영 띠 + 왼쪽 굵은 검정 세로바 + 굵게."""
    text = f"{number}. {title}" if number else title
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, FONT_TITLE, size=14, bold=True)
    set_para_shading(p, SH_SECTION_BAND)
    set_para_border(p, {"left": {"val": "single", "sz": 24, "color": "000000", "space": 4}})
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.1)


# ── 글머리 항목 ──────────────────────────────────────────────────────────
def build_bullet(doc, level, text):
    """위계별 마커(□○-·) + 들여쓰기. 마커는 맑은 고딕, 본문은 바탕."""
    level = max(1, min(4, int(level)))
    marker = BULLET_MARKERS[level]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(BULLET_INDENT_CM[level])
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    rm = p.add_run(f"{marker} ")
    set_run_font(rm, FONT_TITLE, size=11, bold=(level == 1))
    rt = p.add_run(text)
    set_run_font(rt, FONT_BODY, size=11, bold=(level == 1))


def build_para(doc, text, indent_cm=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    set_run_font(r, FONT_BODY, size=11)


# ── 표 ───────────────────────────────────────────────────────────────────
def build_table(doc, spec):
    """얇은 회색 테두리, 머리글 행 회색·굵게·가운데, 라벨 열 연회색·굵게."""
    header_row = spec.get("header_row")     # list[str] | None
    rows = spec.get("rows", [])             # list[list[str]]
    label_col = spec.get("label_col", False)
    col_widths = spec.get("col_widths")     # list[float] (cm) | None

    n_cols = len(header_row) if header_row else (len(rows[0]) if rows else 1)
    n_rows = (1 if header_row else 0) + len(rows)
    if n_rows == 0:
        return

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    r_idx = 0
    if header_row:
        for c, val in enumerate(header_row):
            cell = table.cell(0, c)
            set_cell_shading(cell, SH_TABLE_HEADER)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(val))
            set_run_font(run, FONT_TITLE, size=10.5, bold=True)
        r_idx = 1

    for ri, row in enumerate(rows):
        for c in range(n_cols):
            val = row[c] if c < len(row) else ""
            cell = table.cell(r_idx + ri, c)
            para = cell.paragraphs[0]
            is_label = label_col and c == 0
            if is_label:
                set_cell_shading(cell, SH_LABEL_COL)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(str(val))
                set_run_font(run, FONT_TITLE, size=10.5, bold=True)
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(str(val))
                set_run_font(run, FONT_BODY, size=10.5)

    if col_widths:
        for c, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[c].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── 붙임 / 끝. ───────────────────────────────────────────────────────────
def build_attachment(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    set_run_font(r, FONT_BODY, size=11)


def build_end_mark(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("끝.")
    set_run_font(r, FONT_BODY, size=11)
    p.paragraph_format.space_before = Pt(10)


# ── 쪽번호 (하단 가운데 '- N -') ─────────────────────────────────────────
def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("- ")
    set_run_font(r1, FONT_BODY, size=10)
    r2 = p.add_run()
    set_run_font(r2, FONT_BODY, size=10)
    add_field(r2, "PAGE")
    r3 = p.add_run(" -")
    set_run_font(r3, FONT_BODY, size=10)


# ── 메인 ─────────────────────────────────────────────────────────────────
def render(spec):
    doc = Document()

    # 기본 여백 (A4)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 문서 기본 폰트
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = CLR_BLACK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_BODY)

    build_header(doc, spec.get("header", {}))

    for el in spec.get("body", []):
        t = el.get("type")
        if t == "section":
            build_section_heading(doc, el.get("number", ""), el.get("title", ""))
        elif t == "bullet":
            build_bullet(doc, el.get("level", 1), el.get("text", ""))
        elif t == "para":
            build_para(doc, el.get("text", ""), el.get("indent_cm", 0.0))
        elif t == "table":
            build_table(doc, el)
        elif t == "attachment":
            build_attachment(doc, el.get("text", ""))
        elif t == "end":
            build_end_mark(doc)
        elif t == "spacer":
            doc.add_paragraph()
        else:
            sys.stderr.write(f"[경고] 알 수 없는 요소 type: {t}\n")

    if spec.get("page_number", True):
        add_page_number(doc)

    return doc


def main():
    if len(sys.argv) < 2:
        sys.stderr.write(__doc__)
        sys.exit(1)

    src = sys.argv[1]
    if src == "-":
        spec = json.load(sys.stdin)
    else:
        with open(src, "r", encoding="utf-8") as f:
            spec = json.load(f)

    out = None
    if len(sys.argv) >= 3:
        out = sys.argv[2]
    else:
        out = spec.get("meta", {}).get("output")
    if not out:
        title = spec.get("header", {}).get("title", "정부공문서")
        safe = "".join(c for c in title if c not in '\\/:*?"<>|').strip() or "정부공문서"
        out = f"{safe}.docx"

    doc = render(spec)
    doc.save(out)
    print(f"생성 완료: {os.path.abspath(out)}")


if __name__ == "__main__":
    main()
